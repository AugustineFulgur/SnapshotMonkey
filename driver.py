import pyautogui
import os
import time
import argparse
import cmds
from cmds import *
from docx import Document
import pkgutil
import importlib
import re
import signal
from docx.shared import Cm, Inches

DESCRIPTION='''
by AugustTheodor CatXicure\n
[猴子的命也是命！ Monkey lives matter!]\n
自动化猴子脚本 内网漏洞验证工具\n
目前支持 海康威视命令执行、FTP、SMBEXEC\n
如果要增加验证漏洞出POC 请参照cmds模块下的文件书写 文件会自动加载。\n
使用 python driver.py -t fscan扫描文件\n
'''

SNAPSHOT_MONKEY='''
 ____  _      ____  ____  ____  _     ____  _____    _      ____  _      _  __ ________  _
/ ___\/ \  /|/  _ \/  __\/ ___\/ \ /|/  _ \/__ __\  / \__/|/  _ \/ \  /|/ |/ //  __/\  \//
|    \| |\ ||| / \||  \/||    \| |_||| / \|  / \    | |\/||| / \|| |\ |||   / |  \   \  / 
\___ || | \||| |-|||  __/\___ || | ||| \_/|  | |    | |  ||| \_/|| | \|||   \ |  /_  / /  
\____/\_/  \|\_/ \|\_/   \____/\_/ \|\____/  \_/    \_/  \|\____/\_/  \|\_|\_\\____\/_/   
                                                                                           '''

def get_modules():
    modules_name={}
    for filefiner, name, ispkg in pkgutil.iter_modules(cmds.__path__):
        if not ispkg :
            modules_name[importlib.import_module("cmds."+name).__Aname__()]=importlib.import_module("cmds."+name) #如果标志方法在文件中，加入这个模块（话说python是不是有抽象类 不管了）
            #所以方法中必须带有__Aexec__和__Aname__ 前者执行自动化操作 后者对应FSCAN标志
    print("加载EXECS成功，共{}条。".format(len(modules_name.keys())))
    return modules_name

from win32api import GetMonitorInfo, MonitorFromPoint
def get_workarea():
    monitor_info = GetMonitorInfo(MonitorFromPoint((0,0)))
    work_area = monitor_info.get("Work")
    return work_area

def __main__():
    print(SNAPSHOT_MONKEY) #hzzmnl
    emodules=get_modules() #读取模块
    parser=argparse.ArgumentParser(description=DESCRIPTION)
    parser.add_argument("-t",type=str,help="fscan文件",required=True)
    parser.add_argument("-name",type=str,help="输出doc文件名",required=True)
    argv=parser.parse_args()
    reg_blast=re.compile("([^\:\]]*)\:([0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3}\.[0-9]{1,3})\:[0-9]{2,5}\:(.*)")#素一款正则表达式大师 虽然是多余的正则表达式大师
    #爆破与弱口令、未授权、匿名（合并一起写方便）
    reg_poc=re.compile("[\S]* ([\S]*) ([\S]*)")#一款比较没什么用的正则表达式
    doc=Document() #打开新Docx
    with open(argv.t) as fsc:
       deal_list=fsc.read().split("\n")
       for i in deal_list:
           if i.startswith("[+]"): #如果带有[+]
                p=doc.add_paragraph() #添加段落
                p.add_run(i+"\n") #带有[+]的都写进报告
                try:
                    group=re.search(reg_blast,i)
                    #Aexec的参数和group匹配到的一样[+] [漏洞名] [IP][端口] [用户名 密码]|[匿名/未授权]
                    if group.group(1) in emodules.keys():
                        #匹配到对应的执行方式
                        print("匹配到{},IP{}".format(group.group(1),group.group(2)))
                        pic_route=emodules[group.group(1).strip()].__Aexec__([group.group(2),group.group(3),group.group(4)]) #传入的是 IP 端口 用户名密码/未授权/匿名
                        doc.add_picture(pic_route,width=Cm(15))
                        os.remove(pic_route)
                except:
                    #那就不是上面这种 继续
                    try:
                        group=re.search(reg_poc,i)
                        #[URL] [POC名]
                        if group.group(2).strip() in emodules.keys():
                            print("匹配到{},链接{}".format(group.group(2),group.group(1)))
                            route=emodules[group.group(2).strip()].__Aexec__([group.group(1)])
                            doc.add_picture(route,width=Cm(15))
                            os.remove(pic_route) #删除图片
                    except:
                        continue
    doc.save(argv.name) #保存docx
                    
if __name__=="__main__":
    __main__()